import streamlit as st
import pandas as pd
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
import base64
import datetime

# Set page configuration
st.set_page_config(
    page_title="Biology Lab Room Chart Generator",
    page_icon="🧪",
    layout="wide"
)

# Title and description
st.title("🧪 Biology Lab Room Chart Generator")
st.markdown("Upload your CSV or Excel file to generate a Room Use Chart as a Word document. Specify the semester and year on the left sidebar.")

# Helper functions
def parse_time(time_val):
    if pd.isna(time_val) or time_val == '':
        return None
    
    if isinstance(time_val, datetime.time):
        return time_val.hour * 60 + time_val.minute

    if isinstance(time_val, str):
        try:
            time, period = time_val.strip().split(' ')
            hours, minutes = map(int, time.split(':'))
            if period.upper() == 'PM' and hours != 12:
                hours += 12
            if period.upper() == 'AM' and hours == 12:
                hours = 0
            return hours * 60 + minutes
        except:
            return None
    return None

def extract_room_number(location):
    if pd.isna(location) or location == '':
        return None
    match = re.findall(r'\d+', str(location))
    return int(match[-1]) if match else None

def extract_last_name(instructor):
    if pd.isna(instructor) or instructor == '':
        return ''
    if 'Wilnelia Recart Gonzalez' in instructor:
        return 'RECART'
    return instructor.strip().split(' ')[-1].upper()

def format_time(time_val):
    if pd.isna(time_val) or time_val == '':
        return ''
    
    if isinstance(time_val, datetime.time):
        formatted_time = time_val.strftime('%I:%M')
        if formatted_time.startswith('0'):
            return formatted_time[1:]
        return formatted_time

    if isinstance(time_val, str):
        return re.sub(r'\s*(AM|PM)', '', time_val, flags=re.IGNORECASE)

    return str(time_val)

def abbreviate_title(title):
    if pd.isna(title) or title == '':
        return ''
    abbreviations = {
        'Anatomy & Physiology': 'A & P',
        'Bioenergetics and Systems': 'Bioenergetics',
        'Genomes and Evolution': 'Genome Evol',
        'Medical Microbiology': 'Med Micro',
        'Earth/Life Sci for Educators': 'Life Sci Ed',
        'Biostatistics': 'Biostats',
        'Biology Capstone Seminar': 'Capstone',
        'Insect Biology': 'Insect Bio',
        'Science in the Public Domain': 'SCI Pub Dom',
        'Ecological Community:San Diego': 'Ecol Comm',
        'Research Methods': 'Res Meth',
        'Cell Physiology': 'Cell Phys',
        'Vertebrate Physiology': 'Vert Phys',
        'Microbiology': 'Micro',
        'Research Project': 'Res Proj',
        'Techniques: Molecular Biology': 'Molec Tech',
        'Comp. Anat. of Vertebrates': 'Comp An Vert',
        'Invertebrate Zoology': 'Invert Zoo',
        'Peoples, Plagues and Microbes': 'Ppl Plag Micro',
        'Ecol Evol Infectious Disease': 'EEID',
        'Immunology': 'Immuno',
        'Laboratory': '',
        'Lab': ''
    }
    for full, abbrev in abbreviations.items():
        title = re.sub(full, abbrev, title, flags=re.IGNORECASE)
    return title

def expand_days(days_str):
    if pd.isna(days_str) or days_str == '':
        return []
    mapping = {'M': 'Mon', 'T': 'Tue', 'W': 'Wed', 'R': 'Thu', 'F': 'Fri', 'S': 'Sat', 'U': 'Sun'}
    return [mapping[c] for c in str(days_str) if c in mapping]

def is_before_noon(time_val):
    if pd.isna(time_val) or time_val == '':
        return False

    if isinstance(time_val, datetime.time):
        return time_val.hour < 12

    if isinstance(time_val, str):
        try:
            time_parts = time_val.strip().split(' ')
            if len(time_parts) < 2: return False
            time_part = time_parts[0]
            period = time_parts[1].upper()
            hours, minutes = map(int, time_part.split(':'))
            if period == 'PM' and hours != 12:
                hours += 12
            elif period == 'AM' and hours == 12:
                hours = 0
            return hours < 12
        except:
            return False
    return False

def process_csv_and_generate_doc(uploaded_file, target_rooms, semester, year):
    """Process the CSV or Excel file and generate the Word document"""
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, header=[0, 1])
        elif uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file, header=[0, 1])
        else:
            st.error("Unsupported file type.")
            return None, 0
        
        df.columns = [f"{a} {b}".strip() if not pd.isna(a) and not pd.isna(b) else (a or b) for a, b in df.columns]
        
        df[['Course Number:', 'Title: Unnamed: 3_level_1', 'Instructors: Unnamed: 14_level_1']] = df[['Course Number:', 'Title: Unnamed: 3_level_1', 'Instructors: Unnamed: 14_level_1']].ffill()
        df = df[df['Seats Remaining:'] != 'CLOSED']
        
        location_col = 'Location: Unnamed: 15_level_1'
        course_col = 'Course Number:'
        title_col = 'Title: Unnamed: 3_level_1'
        days_col = 'Days: Unnamed: 6_level_1'
        begin_col = 'Begin Time:'
        end_col = 'End Time:'
        instructor_col = 'Instructors: Unnamed: 14_level_1'
        
        entries = []
        
        # This new try/except block will catch the failing row
        try:
            for index, row in df.iterrows():
                room = extract_room_number(row.get(location_col))
                if room not in target_rooms:
                    continue
                days = expand_days(row.get(days_col))
                for day in days:
                    entries.append({
                        'Day': day,
                        'Room': room,
                        'Course_Number': str(row.get(course_col)).replace('BIOL', 'BIO'),
                        'Title': abbreviate_title(row.get(title_col)),
                        'Begin_Time': format_time(row.get(begin_col)),
                        'End_Time': format_time(row.get(end_col)),
                        'Instructors': extract_last_name(row.get(instructor_col)),
                        'Begin_Time_Parsed': parse_time(row.get(begin_col)),
                        'Begin_Time_Original': row.get(begin_col)
                    })
        except Exception as e:
            st.error(f"### 🚨 An error occurred while processing the data.")
            st.error(f"**Error message:** `{str(e)}`")
            st.error(f"This likely happened on **row {index + 2}** of your Excel file.")
            st.warning("**Dumping data from the row that caused the error:**")
            # Convert the row to a dictionary for clean printing, handling potential datetime objects
            st.json({k: str(v) for k, v in row.to_dict().items()})
            return None, 0

        day_order = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
        entries.sort(key=lambda x: (day_order.index(x['Day']), x['Room'], x['Begin_Time_Parsed'] or 0))
        
        rooms_sorted = sorted(target_rooms)
        days_present = sorted(set(e['Day'] for e in entries), key=lambda d: day_order.index(d))
        chart_rows = []
        
        for day in days_present:
            row = {'Day': day}
            for room in rooms_sorted:
                classes = [e for e in entries if e['Day'] == day and e['Room'] == room]
                row[f'ST{room}'] = classes if classes else ''
            chart_rows.append(row)
        
        chart_df = pd.DataFrame(chart_rows)
        
        doc = Document()
        section = doc.sections[0]
        section.page_height, section.page_width = Inches(8.5), Inches(11)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
        
        title_text = f"{semester} {year} Room Use Chart for the Biology Laboratories"
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Font settings... (rest of the document generation is the same)
        
        # (The rest of your Word document generation code is unchanged)
        # Create table
        table = doc.add_table(rows=1, cols=len(chart_df.columns))
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set table width
        table_width = Inches(10)
        table.width = int(table_width)
        
        # Calculate column widths
        day_col_width = Inches(1.2)
        room_col_width = (table_width - day_col_width) / (len(chart_df.columns) - 1)
        
        # Set column widths
        for i, col in enumerate(table.columns):
            if i == 0:
                col.width = int(day_col_width)
            else:
                col.width = int(room_col_width)
        
        # Format header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(chart_df.columns):
            cell = hdr_cells[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if col_name == 'Day':
                cell.paragraphs[0].clear()
                run1 = cell.paragraphs[0].add_run("B = Morning")
                run1.font.name = 'Times New Roman'
                run1.bold = True
                run1.font.size = Pt(12)
                run1.font.color.rgb = RGBColor(0, 0, 255)
                
                run2 = cell.paragraphs[0].add_run("\n")
                run2.font.name = 'Times New Roman'
                
                run3 = cell.paragraphs[0].add_run("G = Afternoon")
                run3.font.name = 'Times New Roman'
                run3.bold = True
                run3.font.size = Pt(12)
                run3.font.color.rgb = RGBColor(0, 128, 0)
            elif col_name.startswith('ST'):
                run = cell.paragraphs[0].add_run(str(col_name))
                run.font.name = 'Times New Roman'
                run.bold = True
                run.font.size = Pt(18)
            else:
                run = cell.paragraphs[0].add_run(str(col_name))
                run.font.name = 'Times New Roman'
                run.bold = True
                run.font.size = Pt(12)
        
        # Add data rows
        for _, row in chart_df.iterrows():
            row_cells = table.add_row().cells
            
            for j, val in enumerate(row):
                para = row_cells[j].paragraphs[0]
                col_name = chart_df.columns[j]
                
                if col_name == 'Day':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(str(val))
                    run.font.name = 'Times New Roman'
                    run.bold = True
                    run.font.size = Pt(20)
                elif col_name.startswith('ST'):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if isinstance(val, list) and val:
                        morning_classes = []
                        afternoon_classes = []
                        
                        for e in val:
                            if is_before_noon(e['Begin_Time_Original']):
                                morning_classes.append(e)
                            else:
                                afternoon_classes.append(e)
                        
                        only_afternoon = len(morning_classes) == 0 and len(afternoon_classes) > 0
                        
                        if only_afternoon:
                            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                        else:
                            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                        
                        run = para.add_run("\n")
                        run.font.name = 'Times New Roman'
                        
                        for idx, e in enumerate(morning_classes):
                            if idx > 0:
                                run = para.add_run("\n\n")
                                run.font.name = 'Times New Roman'
                            
                            class_text = f"{e['Begin_Time']}-{e['End_Time']}\n{e['Course_Number']}\n{e['Title']}\n{e['Instructors']}"
                            run = para.add_run(class_text)
                            run.font.name = 'Times New Roman'
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 255)
                        
                        if morning_classes and afternoon_classes:
                            run = para.add_run("\n\n")
                            run.font.name = 'Times New Roman'
                        
                        for idx, e in enumerate(afternoon_classes):
                            if idx > 0:
                                run = para.add_run("\n\n")
                                run.font.name = 'Times New Roman'
                            
                            class_text = f"{e['Begin_Time']}-{e['End_Time']}\n{e['Course_Number']}\n{e['Title']}\n{e['Instructors']}"
                            run = para.add_run(class_text)
                            run.font.name = 'Times New Roman'
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        run = para.add_run('')
                        run.font.name = 'Times New Roman'

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer, len(entries)
    
    except Exception as e:
        st.error(f"A critical error occurred before data processing could start: {str(e)}")
        return None, 0

# (The rest of your Streamlit UI code is unchanged)
# Streamlit UI
st.sidebar.header("Settings")

# Semester and Year selection
st.sidebar.subheader("Semester & Year")
col1, col2 = st.sidebar.columns(2)
with col1:
    semester = st.selectbox(
        "Semester:",
        options=["Spring", "Fall"],
        index=0
    )
with col2:
    year = st.number_input(
        "Year:",
        min_value=2020,
        max_value=2100,
        value=2025,
        step=1
    )

# Add validation
if year < 2020 or year > 2100:
    st.warning("Please enter a valid year between 2020 and 2100")

# Set target rooms (no user selection needed)
target_rooms = [225, 227, 229, 242, 325, 327, 330, 429]

# File upload
uploaded_file = st.file_uploader(
    "Upload your CSV or Excel file",
    type=['csv', 'xlsx'],
    help="Upload a CSV or Excel file with course schedule data"
)

if uploaded_file is not None:
    st.success("File uploaded successfully!")
    
    # Show file info
    st.info(f"File name: {uploaded_file.name}")
    
    # Process button
    if st.button("Generate Room Chart", type="primary"):
        if year < 2020 or year > 2100:
            st.error("Please enter a valid year between 2020 and 2100 before generating the chart.")
        else:
            with st.spinner("Processing file and generating Word document..."):
                doc_buffer, num_entries = process_csv_and_generate_doc(uploaded_file, target_rooms, semester, year)
                
                if doc_buffer is not None:
                    st.success(f"Document generated successfully! Found {num_entries} class entries.")
                    
                    # Download button with dynamic filename
                    filename = f"{semester}_{year}_room_use_chart.docx"
                    st.download_button(
                        label="📥 Download Room Chart (Word Document)",
                        data=doc_buffer,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Failed to generate document. Please check your file format and the error messages above.")

# Instructions
with st.expander("📋 Instructions"):
    st.markdown("""
    ### How to use this tool:
    1.  Upload your CSV or Excel file above.
    2.  Select the semester and year in the sidebar.
    3.  Click "Generate Room Chart".
    4.  Download the Word document.
    5.  **Review & Finalize**: Always cross-check the generated document with the original schedule and make any necessary final edits.
    
    ### How to create the input file:
    -   Go to the class schedule website: https://usdssb.sandiego.edu/prod/usd_course_query_faculty.p_start 
    -   Select the desired semester and "Biology" as the department, then click Submit.
    -   Copy all data from the "CRN:" column to the "Location:" column.
    -   Paste the data into Excel and save as a `.csv` or `.xlsx` file.
    
    ### Output:
    -   A formatted Room Use Chart in a Word document.
    """)
